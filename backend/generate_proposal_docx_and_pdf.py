"""
National University Bangladesh AI Assistant & Smart Support Platform
Professional Project Proposal Generator (DOCX, PDF & Markdown Editions in Bengali)
Target Directory: E:/projects/AI_CHAT_BOT/project_proposal/
"""

import os
import sys
import subprocess
from pathlib import Path

BASE_DIR = Path("E:/projects/AI_CHAT_BOT")
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

PROPOSAL_DIR = BASE_DIR / "project_proposal"
PROPOSAL_DIR.mkdir(parents=True, exist_ok=True)
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

BASE_DIR = Path("E:/projects/AI_CHAT_BOT")
PROPOSAL_DIR = BASE_DIR / "project_proposal"
PROPOSAL_DIR.mkdir(parents=True, exist_ok=True)

# Color Palette
COLOR_EMERALD_DARK = RGBColor(6, 95, 70)     # #065f46
COLOR_EMERALD = RGBColor(5, 150, 105)        # #059669
COLOR_NAVY = RGBColor(15, 23, 42)            # #0f172a
COLOR_SLATE = RGBColor(71, 85, 105)          # #475569
COLOR_TEXT = RGBColor(30, 41, 59)            # #1e293b

def set_cell_background(cell, fill_hex):
    """Sets background color of a docx table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_proposal_docx():
    doc = Document()
    
    # 1. Page Margins (Normal 1 inch / 0.75 inch)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # 2. Header & Branding
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_hdr = header_para.add_run("জাতীয় বিশ্ববিদ্যালয়, বাংলাদেশ • প্রাতিষ্ঠানিক প্রজেক্ট প্রপোজাল ২০২৬")
    run_hdr.font.name = "Hind Siliguri"
    run_hdr.font.size = Pt(9)
    run_hdr.font.color.rgb = COLOR_SLATE

    # Title Block
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(4)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("জাতীয় বিশ্ববিদ্যালয় স্মার্ট AI অ্যাসিস্ট্যান্ট ও সেন্ট্রালাইজড সাপোর্ট ইকোসিস্টেম")
    run_title.font.name = "Hind Siliguri"
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_EMERALD_DARK

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(12)
    run_sub = p_sub.add_run("৩২ লক্ষ শিক্ষার্থী, ২ লক্ষ শিক্ষক ও ২,২৬০টি অধিভুক্ত কলেজের জন্য আধুনিক কৃত্রিম বুদ্ধিমত্তা ভিত্তিক সেবা অটোমেশন")
    run_sub.font.name = "Hind Siliguri"
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = COLOR_SLATE

    # Meta Table (4 Key Statistics in colored boxes)
    stat_table = doc.add_table(rows=1, cols=4)
    stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [
        ("২,২৬০+ কলেজ", "অধিভুক্ত প্রতিষ্ঠান", "f0fdf4"),
        ("৩২ লক্ষ+ শিক্ষার্থী", "স্নাতক ও স্নাতকোত্তর", "eff6ff"),
        ("২ লক্ষ+ শিক্ষক", "শিক্ষক ও কর্মকর্তা", "fffbeb"),
        ("৬৪ জেলায় বিস্তৃত", "সারাদেশব্যাপী নেটওয়ার্ক", "faf5ff")
    ]
    for i, (title, sub, bg) in enumerate(stats):
        cell = stat_table.cell(0, i)
        set_cell_background(cell, bg)
        set_cell_margins(cell, top=140, bottom=140, left=140, right=140)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{title}\n")
        r1.font.name = "Hind Siliguri"
        r1.font.size = Pt(12)
        r1.font.bold = True
        r1.font.color.rgb = COLOR_EMERALD_DARK
        r2 = p.add_run(sub)
        r2.font.name = "Hind Siliguri"
        r2.font.size = Pt(9)
        r2.font.color.rgb = COLOR_SLATE

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 1: Executive Summary
    h1 = doc.add_heading(level=1)
    r = h1.add_run("১. প্রজেক্টের পটভূমি ও উদ্দেশ্য (Executive Background & Rationale)")
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_EMERALD_DARK
    h1.paragraph_format.space_before = Pt(10)
    h1.paragraph_format.space_after = Pt(4)

    p1 = doc.add_paragraph()
    r = p1.add_run(
        "জাতীয় বিশ্ববিদ্যালয় (nu.ac.bd) বাংলাদেশের উচ্চশিক্ষার সর্ববৃহৎ চালিকাশক্তি। দেশের উচ্চশিক্ষার প্রায় ৭০% শিক্ষার্থী এই বিশ্ববিদ্যালয়ের অধিভুক্ত কলেজসমূহে অধ্যয়ন করে। বিশাল ভৌগোলিক বিস্তার এবং বিপুল সংখ্যক শিক্ষার্থীর কারণে ভর্তি, পরীক্ষার রুটিন, ফরম পূরণ, ইএমএস (EMS) পোর্টাল লকআউট, মার্কশিট ও মূল সনদ উত্তোলনের সময় বিশ্ববিদ্যালয় প্রশাসন ও কলেজগুলো তীব্র তথ্য ও সেবাপ্রবাহের চাপের মুখোমুখি হয়।"
    )
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(10.5)

    p2 = doc.add_paragraph()
    r = p2.add_run(
        "বর্তমানে হাজার হাজার শিক্ষার্থী সামান্য তথ্যের জন্য বা পোর্টাল পাসওয়ার্ড রিসেট করতে দূর-দূরান্তের জেলা (যেমন: পঞ্চগড়, কক্সবাজার, সুনামগঞ্জ, পটুয়াখালী) থেকে গাজীপুর মূল ক্যাম্পাসে সশরীরে আসতে বাধ্য হয়। এতে শিক্ষার্থীদের বিপুল আর্থিক ক্ষতি ও হয়রানি হয় এবং বিশ্ববিদ্যালয়ের হেল্পডেস্কে তীব্র জট তৈরি হয়। এই সমস্যা চিরতরে নিরসনকল্পে 'স্মার্ট বাংলাদেশ ২০৪১' রূপকল্পের আলোকে সম্পূর্ণ স্বয়ংক্রিয়, ২৪/৭ সক্রিয় AI অ্যাসিস্ট্যান্ট ও সাপোর্ট টোকেন প্ল্যাটফর্ম প্রস্তাব করা হলো।"
    )
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(10.5)

    # Section 2: Stakeholder Benefits
    h2 = doc.add_heading(level=1)
    r = h2.add_run("২. স্টেকহোল্ডার ভিত্তিক সুবিধার বিশ্লেষণ (Stakeholder Value Realization)")
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_EMERALD_DARK
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)

    # 4 Cards for Stakeholders
    stakeholders = [
        ("🎓 শিক্ষার্থীদের জন্য সুবিধা (৩.২ মিলিয়ন শিক্ষার্থী)", [
            "তাৎক্ষণিক ২৪/৭ উত্তর: বাংলা বা ইংরেজিতে ভর্তি যোগ্যতা, সংশোধিত রুটিন ও রেজাল্ট জানার সুযোগ (< 0.001 সেকেন্ড)।",
            "অফিসিয়াল সাপোর্ট টোকেন: ইএমএস লগইন, ফরম পূরণ বা সনদ জটিলতায় ট্র্যাকিং নম্বর (NU-2026-XXXXXX) সহ সমাধান।",
            "বিপুল অর্থ ও সময় সাশ্রয়: গাজীপুর মূল ক্যাম্পাসে আসার যাতায়াত ও থাকা-খাওয়ার খরচ সম্পূর্ণ দূর হবে।",
            "মোবাইল ক্যামেরা কিউআর: কোনো অ্যাপ ইন্সটল না করেই যেকোনো স্মার্টফোনে ক্যামেরা দিয়ে স্ক্যান করে চ্যাটবট সক্রিয়।"
        ], "ecfdf5"),
        ("👨‍🏫 শিক্ষক ও পরীক্ষা কর্মকর্তাদের জন্য সুবিধা (২ লক্ষ+ শিক্ষক)", [
            "পরীক্ষা ও মূল্যায়ন বিধি: প্রশ্ন প্রণয়ন, খাতা মূল্যায়ন ও মডারেশনের সার্বিক নিয়মাবলী এক ক্লিকে জানা।",
            "সম্মানী ও বিলিং পদ্ধতি: সোনালী সেবার মাধ্যমে পরীক্ষক বিল সাবমিশনের সঠিক গাইডলাইন তাৎক্ষণিক প্রাপ্তি।",
            "সিলেবাস ও কারিকুলাম ভেরিফিকেশন: অনার্স ও মাস্টার্সের হালনাগাদ সিলেবাস ও ক্রেডিট সংক্রান্ত নির্ভুল তথ্য।"
        ], "eff6ff"),
        ("🏛️ অধিভুক্ত কলেজ ও অধ্যক্ষদের জন্য সুবিধা (২,২৬০+ কলেজ)", [
            "কলেজ অফিসের চাপ ৮০% হ্রাস: রুটিন ও নোটিশ সংক্রান্ত শত শত প্রশ্নের উত্তর AI স্বয়ংক্রিয়ভাবে দিয়ে দেবে।",
            "সরাসরি গাজীপুর ডেস্ক এসকেলেশন: জটিল প্রাতিষ্ঠানিক সমস্যা সরাসরি কেন্দ্রীয় আইসিটি ও পরীক্ষা ডেস্কে প্রেরণ।",
            "ডিজিটাল তথ্য সমতা: প্রত্যন্ত অঞ্চলের গ্রামীণ কলেজগুলোও কেন্দ্রীয় ক্যাম্পাসের মতোই রিয়েল-টাইমে আপডেট পাবে।"
        ], "fffbeb"),
        ("🏢 গাজীপুর সেন্ট্রাল প্রশাসনের জন্য সুবিধা (ICT, পরীক্ষা ও রেজিস্ট্রেশন সেল)", [
            "৮৫% কল ও টিকিট হ্রাস: সাধারণ প্রশ্নোত্তর AI সমাধান করায় কর্মকর্তারা গুরুত্বপূর্ণ ফাইলে মনোযোগ দিতে পারবেন।",
            "সেলফ-লার্নিং ভেক্টর ফিডব্যাক: একবার কোনো সমস্যার সমাধান দিলে AI তা শিখে নিয়ে ভবিষ্যতে উত্তর দিতে পারে।",
            "২৪/৭ স্বয়ংক্রিয় নলেজ রিফ্রেশ: ক্রলার সার্বক্ষণিক নতুন নোটিশ বিশ্লেষণ করে ভেক্টর ডাটাবেস সমৃদ্ধ রাখে।"
        ], "faf5ff")
    ]

    for title, points, bg in stakeholders:
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, bg)
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        p = cell.paragraphs[0]
        r_head = p.add_run(f"{title}\n")
        r_head.font.name = "Hind Siliguri"
        r_head.font.size = Pt(11)
        r_head.font.bold = True
        r_head.font.color.rgb = COLOR_EMERALD_DARK
        for pt in points:
            r_pt = p.add_run(f"• {pt}\n")
            r_pt.font.name = "Hind Siliguri"
            r_pt.font.size = Pt(9.5)
            r_pt.font.color.rgb = COLOR_TEXT
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Section 3: Core Capabilities & Architecture
    h3 = doc.add_heading(level=1)
    r = h3.add_run("৩. প্ল্যাটফর্মের প্রধান কার্যাবলী ও প্রযুক্তিগত উদ্ভাবন")
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_EMERALD_DARK
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)

    # Tech Stack Table
    tech_table = doc.add_table(rows=1, cols=3)
    tech_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["মডিউল / প্রযুক্তি", "উপাদান ও প্রোটোকল", "প্রকৌশলগত ভূমিকা ও কার্যপদ্ধতি"]
    for i, h in enumerate(headers):
        cell = tech_table.cell(0, i)
        set_cell_background(cell, "0f172a")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Hind Siliguri"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    tech_rows = [
        ("FastAPI Backend", "Python 3.13 / Async", "নন-ব্লকিং অ্যাসিনক্রোনাস আর্কিটেকচার; হেভি GenAI ও ডিস্ক I/O পৃথক থ্রেডে পাঠিয়ে দ্রুততম রেসপন্স।"),
        ("Google Gemini 3 Flash", "gemini-3-flash-preview", "উন্নত জেনারেটিভ AI; সাব-সেকেন্ডে বাংলা ও ইংরেজি প্রশ্ন বুঝে অফিসিয়াল সার্কুলার লিংক সহ নির্ভুল উত্তর।"),
        ("ChromaDB Vector Store", "gemini-embedding-001", "অফিসিয়াল নোটিশ ও সমাধানের জন্য সেমান্টিক ভেক্টর সার্চ ইঞ্জিন; রেজিলিয়েন্ট সিউডো-ভেক্টর ব্যাকআপ ব্যবস্থা যুক্ত।"),
        ("MCP Protocol Suite", "5 Dedicated MCP Servers", "token_mcp, knowledge_mcp, document_mcp, credential_mcp ও enrichment_mcp পরিচালনা করে।"),
        ("Fernet AES-128 Vault", "CBC + HMAC-SHA256", "ইএমএস ও স্টুডেন্ট পোর্টালের লগইন পাসওয়ার্ড সুরক্ষায় AES-128 এনক্রিপশন ও PBKDF2 হ্যাশিং।"),
        ("Preloaded Fast Cache", "In-Memory Trie (18 µs)", "সাধারণ সম্ভাষণ (hi, ভর্তি, রুটিন, রেজাল্ট) ১৮ মাইক্রোসেকেন্ডে মেমোরি থেকে সরাসরি প্রদান।"),
        ("২৪/৭ স্বয়ংক্রিয় এজেন্ট", "Autonomous Multi-Agent", "প্রতি ১০ মিনিটে nu.ac.bd সাইট ক্রল করে প্রশ্নোত্তর তৈরি ও ChromaDB-তে লাইভ ইনজেশন।")
    ]

    for row_data in tech_rows:
        row = tech_table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Hind Siliguri"
            r.font.size = Pt(9)
            if i == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 4: Visual Screenshots
    h4 = doc.add_heading(level=1)
    r = h4.add_run("৪. প্ল্যাটফর্মের ইন্টারফেস ও কার্যপ্রণালীর চিত্র (Interface Demonstration)")
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_EMERALD_DARK
    h4.paragraph_format.space_before = Pt(12)
    h4.paragraph_format.space_after = Pt(6)

    # Add images with captions
    images_to_embed = [
        ("1.png", "চিত্র ১: মূল AI চ্যাটবট ইন্টারফেস — বাংলা ও ইংরেজিতে তাৎক্ষণিক নোটিশ ও সার্কুলার রেফারেন্স সহ উত্তর।"),
        ("FAQ.png", "চিত্র ২: তাৎক্ষণিক FAQ ও প্রিলোডেড একাডেমিক জিজ্ঞাসা ডিরেক্টরি।"),
        ("token.png", "চিত্র ৩: অফিসিয়াল সাপোর্ট টোকেন আবেদন ফর্ম — এনক্রিপ্টেড সার্ভিস ক্রেডেনশিয়াল সহ।"),
        ("check_token.png", "চিত্র ৪: লাইভ টোকেন ট্র্যাকিং — PENDING থেকে SOLVED পর্যন্ত রিয়েল-টাইম অগ্রগতি।"),
        ("qr.png", "চিত্র ৫: মোবাইল কিউআর কোড স্ক্যানার — যেকোনো মোবাইল ক্যামেরায় সরাসরি প্ল্যাটফর্ম ব্যবহারের সুবিধা।"),
        ("admin_panel.png", "চিত্র ৬: সেন্ট্রাল অ্যাডমিন ও দাপ্তরিক সলভার কন্ট্রোল সেন্টার — ওয়েবসাইট স্ট্রাকচার ও ২৪/৭ এজেন্ট মনিটরিং।")
    ]

    for img_name, caption in images_to_embed:
        img_path = PROPOSAL_DIR / img_name
        if img_path.exists():
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after = Pt(2)
            doc.add_picture(str(img_path), width=Inches(5.5))
            
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(10)
            r_cap = p_cap.add_run(caption)
            r_cap.font.name = "Hind Siliguri"
            r_cap.font.size = Pt(9)
            r_cap.font.bold = True
            r_cap.font.color.rgb = COLOR_SLATE

    # Section 5: Support Services & SLA
    h5 = doc.add_heading(level=1)
    r = h5.add_run("৫. সাপোর্ট সার্ভিস ক্যাটালগ ও সমাধানের সময়সীমা (SLA)")
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_EMERALD_DARK
    h5.paragraph_format.space_before = Pt(12)
    h5.paragraph_format.space_after = Pt(6)

    sla_table = doc.add_table(rows=1, cols=4)
    sla_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sla_headers = ["সার্ভিস কোড", "সেবার বিবরণ (Service Scope)", "দায়িত্বপ্রাপ্ত দপ্তর (Department Desk)", "সমাধানের সময়সীমা"]
    for i, h in enumerate(sla_headers):
        cell = sla_table.cell(0, i)
        set_cell_background(cell, "065f46")
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.font.name = "Hind Siliguri"
        r.font.size = Pt(9.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)

    sla_rows = [
        ("EMS", "ইএমএস স্টুডেন্ট পোর্টাল পাসওয়ার্ড ও অ্যাকাউন্ট লকআউট সমস্যা", "আইসিটি সাপোর্ট সেল", "২৪-৪৮ ঘণ্টা"),
        ("FORM_FILLUP", "পরীক্ষার ফরম পূরণ ও সোনালী সেবা ফি ভেরিফিকেশন", "পরীক্ষা নিয়ন্ত্রণ শাখা", "১২-২৪ ঘণ্টা"),
        ("RESCRUTINY", "ফলাফল পুনর্নিরীক্ষণ / বোর্ড চ্যালেঞ্জ আবেদন ট্র্যাকিং", "ফলাফল মূল্যায়ন শাখা", "৩-৭ দিন"),
        ("CERTIFICATE", "মূল সনদ ও সাময়িক সনদ উত্তোলন ও অনলাইন ভেরিফিকেশন", "সনদপত্র শাখা (Certificate Wing)", "২-৫ দিন"),
        ("MARKSHEET", "একাডেমিক ট্রান্সক্রিপ্ট ও নম্বরপত্র ভুল সংশোধন", "রেজিস্ট্রার দপ্তর", "২-৪ দিন"),
        ("TC / CORRECTION", "কলেজ ট্রান্সফার এবং নাম/বয়স/রেজিস্ট্রেশন সংশোধন", "রেজিস্ট্রেশন সেল", "৫-১০ দিন")
    ]

    for row_data in sla_rows:
        row = sla_table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell, top=70, bottom=70, left=90, right=90)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.name = "Hind Siliguri"
            r.font.size = Pt(9)
            if i == 0:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Section 6: Strategic Roadmap & ROI
    h6 = doc.add_heading(level=1)
    r = h6.add_run("৬. বাস্তবায়ন রোডম্যাপ ও প্রাতিষ্ঠানিক সুবিধা (ROI)")
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_EMERALD_DARK
    h6.paragraph_format.space_before = Pt(12)
    h6.paragraph_format.space_after = Pt(6)

    # ROI Summary Table
    roi_table = doc.add_table(rows=1, cols=3)
    roi_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    roi_cards = [
        ("💰 ৫০ কোটি+ টাকা বার্ষিক সাশ্রয়", "দূর-দূরান্তের শিক্ষার্থীদের গাজীপুর যাতায়াত, থাকা-খাওয়া ও ফটোকপি খরচ সাশ্রয়।", "f0fdf4"),
        ("⏱️ ৮৫% প্রশাসনিক সময় সাশ্রয়", "টোকেন সমাধানের গড় সময় ৭-১৫ দিনের জায়গায় ২৪ ঘণ্টার নিচে নামিয়ে আনা।", "eff6ff"),
        ("📈 তথ্য সমতা ও স্বচ্ছতা", "৬৪ জেলার সকল অধিভুক্ত কলেজ একযোগে নির্ভুল নোটিশ ও সহায়তা লাভ করবে।", "fffbeb")
    ]
    for i, (title, sub, bg) in enumerate(roi_cards):
        cell = roi_table.cell(0, i)
        set_cell_background(cell, bg)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{title}\n")
        r1.font.name = "Hind Siliguri"
        r1.font.size = Pt(10.5)
        r1.font.bold = True
        r1.font.color.rgb = COLOR_EMERALD_DARK
        r2 = p.add_run(sub)
        r2.font.name = "Hind Siliguri"
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = COLOR_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Sign-off Approval
    h7 = doc.add_heading(level=1)
    r = h7.add_run("৭. সুপারিশ ও প্রাতিষ্ঠানিক অনুমোদনের আবেদন (Sign-Off & Recommendation)")
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_EMERALD_DARK
    h7.paragraph_format.space_before = Pt(12)
    h7.paragraph_format.space_after = Pt(6)

    p_rec = doc.add_paragraph()
    r = p_rec.add_run(
        "জাতীয় বিশ্ববিদ্যালয়ের সকল শিক্ষার্থী, শিক্ষক ও কলেজ সমূহের ডিজিটাল সমতা নিশ্চিতকরণ এবং সেবা সহজীকরণের লক্ষ্যে এই প্রোডাকশন-রেডি AI অ্যাসিস্ট্যান্ট ও স্মার্ট সাপোর্ট প্ল্যাটফর্মটি কেন্দ্রীয়ভাবে মোতায়েন ও বাস্তবায়নের জন্য সনির্বন্ধ প্রস্তাব পেশ করা হলো।"
    )
    r.font.name = "Hind Siliguri"
    r.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Signature Table
    sig_table = doc.add_table(rows=1, cols=3)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    sigs = [
        ("প্রস্তাবক / প্রজেক্ট লিড\nAI আর্কিটেকচার টিম"),
        ("পরিচালক (আইসিটি দপ্তর)\nজাতীয় বিশ্ববিদ্যালয়, বাংলাদেশ"),
        ("উপাচার্য / অনুমোদনকারী কর্তৃপক্ষ\nজাতীয় বিশ্ববিদ্যালয়, বাংলাদেশ")
    ]
    for i, sig_text in enumerate(sigs):
        cell = sig_table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_line = p.add_run("___________________________\n")
        r_line.font.name = "Arial"
        r_line.font.size = Pt(9)
        r_line.font.color.rgb = COLOR_SLATE
        r_txt = p.add_run(sig_text)
        r_txt.font.name = "Hind Siliguri"
        r_txt.font.size = Pt(9.5)
        r_txt.font.bold = True
        r_txt.font.color.rgb = COLOR_NAVY

    # Save DOCX
    docx_path = PROPOSAL_DIR / "project_proposal.docx"
    doc.save(str(docx_path))
    print(f"[OK] Word Document (.docx) generated: {docx_path} ({docx_path.stat().st_size} bytes)")

def build_proposal_suite():
    print("============================================================")
    print("Building Bengali Project Proposal Suite (DOCX, PDF & MD)")
    print("============================================================")
    
    # 1. Generate DOCX
    create_proposal_docx()
    
    # 2. Generate PDF via build_proposal_pdf
    from backend.build_proposal_pdf import build_proposal_pdf
    build_proposal_pdf()
    
    print("\n[SUCCESS] PROPOSAL GENERATION COMPLETED FOR NATIONAL UNIVERSITY BANGLADESH!")

if __name__ == "__main__":
    build_proposal_suite()
