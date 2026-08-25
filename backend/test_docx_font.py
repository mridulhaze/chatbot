import docx
from docx import Document
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor, Inches

doc = Document()

# Set document defaults in styles.xml
styles = doc.styles
normal_style = styles['Normal']
normal_style.font.name = 'Kalpurush'
normal_style.font.size = Pt(11)

rFonts_xml = f'<w:rFonts {nsdecls("w")} w:ascii="Kalpurush" w:hAnsi="Kalpurush" w:cs="Kalpurush" w:eastAsia="Kalpurush"/>'
normal_rPr = normal_style._element.get_or_add_rPr()
normal_rPr.append(parse_xml(rFonts_xml))

p = doc.add_paragraph()
r = p.add_run('জাতীয় বিশ্ববিদ্যালয় AI অ্যাসিস্ট্যান্ট ও স্মার্ট সাপোর্ট প্ল্যাটফর্ম')
rPr = r._r.get_or_add_rPr()
rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="Kalpurush" w:hAnsi="Kalpurush" w:cs="Kalpurush" w:eastAsia="Kalpurush"/>'))
rPr.append(parse_xml(f'<w:lang {nsdecls("w")} w:val="bn-BD" w:bidi="bn-BD"/>'))

doc.save('E:/projects/AI_CHAT_BOT/docs/test_bn_font.docx')
print('Successfully created E:/projects/AI_CHAT_BOT/docs/test_bn_font.docx')
