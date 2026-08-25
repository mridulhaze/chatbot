import re
import io
import csv
import hashlib
import logging
from pathlib import Path
from urllib.parse import urlparse, urljoin, urlunparse, parse_qs, urlencode
from typing import Dict, Any, List, Optional, Tuple

from bs4 import BeautifulSoup
from pypdf import PdfReader

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None

logger = logging.getLogger("NU_CRAWLER_EXTRACTOR")

# Tracking parameters to strip during normalization
STRIP_QUERY_PARAMS = {
    "utm_source", "utm_medium", "utm_campaign", "utm_term", "utm_content",
    "fbclid", "gclid", "sessionid", "phpsessid", "jsessionid", "ref", "source"
}

# Classification rules with weighted priorities
SECTION_RULES = [
    ("Notices", ["notice", "circular", "সংবাদ", "বিজ্ঞপ্তি", "প্রজ্ঞাপন", "recent-notices.php"], 100, "NOTICE"),
    ("Admission", ["admission", "ভর্তি", "app1.nu.edu.bd", "honours", "masters", "degree pass", "professional"], 95, "ADMISSION"),
    ("Examination", ["exam", "পরীক্ষা", "routine", "সময়সূচি", "form fill-up", "ফরম পূরণ", "center", "কেন্দ্র"], 95, "EXAMINATION"),
    ("Results", ["result", "ফলাফল", "results.nu.ac.bd", "rescrutiny", "পুনঃনিরীক্ষণ", "marksheet", "নম্বরপত্র"], 95, "RESULT"),
    ("Forms & Syllabi", ["syllabus", "সিলেবাস", "form", "ফরম", "curriculum", "কারিকুলাম", "download-form"], 90, "FORM_FILLUP"),
    ("Documents", [".pdf", ".docx", ".doc", ".xlsx", ".xls", "/uploads/"], 85, "DOCUMENT"),
    ("Academic", ["academic", "faculty", "অনুষদ", "curriculum", "courses", "বিভাগ", "department"], 70, "ACADEMIC"),
    ("Administration", ["administration", "vc-profile", "pro-vc", "registrar", "treasurer", "কর্মকর্তা", "directory"], 60, "ADMINISTRATION"),
    ("Affiliated Colleges", ["college", "অধিভুক্ত", "affiliated", "college-list"], 50, "COLLEGE"),
    ("General", [], 40, "GENERAL")
]

DATE_PATTERNS = [
    re.compile(r'(\d{1,2})[-/\. ](Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*[-/\. ](\d{4})', re.IGNORECASE),
    re.compile(r'(\d{4})[-/\.](\d{1,2})[-/\.](\d{1,2})'),
    re.compile(r'(\d{1,2})[-/\.](\d{1,2})[-/\.](\d{4})'),
    re.compile(r'(\d{1,2})\s+(জানুয়ারি|ফেব্রুয়ারি|মার্চ|এপ্রিল|মে|জুন|জুলাই|আগস্ট|সেপ্টেম্বর|অক্টোবর|নভেম্বর|ডিসেম্বর)\s+(\d{4})'),
    re.compile(r'(\d{1,2})[-/](\d{1,2})[-/](\d{4})')
]

def normalize_url(url: str, base_url: Optional[str] = None) -> str:
    """Normalizes URL by handling relative paths, stripping fragments and tracking params, standardizing host."""
    if not url:
        return ""
    
    url = url.strip()
    if base_url:
        url = urljoin(base_url, url)
        
    try:
        parsed = urlparse(url)
    except Exception:
        return ""

    if parsed.scheme not in ("http", "https"):
        return ""

    scheme = parsed.scheme.lower()
    netloc = parsed.netloc.lower().split(":")[0]  # strip default port

    # Normalize root host www.nu.ac.bd and nu.ac.bd
    if netloc == "nu.ac.bd":
        netloc = "www.nu.ac.bd"

    path = parsed.path or "/"
    # Normalize duplicate slashes
    path = re.sub(r'/{2,}', '/', path)
    if path.endswith("/index.php") or path.endswith("/index.html"):
        path = path.rsplit("/", 1)[0] + "/"

    # Clean query params
    query = ""
    if parsed.query:
        qs = parse_qs(parsed.query, keep_blank_values=False)
        cleaned_qs = {k: v for k, v in qs.items() if k.lower() not in STRIP_QUERY_PARAMS}
        if cleaned_qs:
            query = urlencode(cleaned_qs, doseq=True)

    normalized = urlunparse((scheme, netloc, path, "", query, ""))
    return normalized

def compute_sha256(content: str or bytes) -> str:
    if isinstance(content, str):
        content = content.encode("utf-8", errors="replace")
    return hashlib.sha256(content).hexdigest()

def clean_extracted_text(text: str) -> str:
    if not text:
        return ""
    # Normalize whitespaces
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()

def detect_language(text: str) -> str:
    # Check for Bengali unicode range
    bengali_chars = len(re.findall(r'[\u0980-\u09FF]', text))
    if bengali_chars > 20 or (len(text) > 0 and (bengali_chars / len(text)) > 0.15):
        return "bn"
    return "en"

def extract_publication_date(text: str) -> Optional[str]:
    for pattern in DATE_PATTERNS:
        match = pattern.search(text)
        if match:
            return match.group(0)
    return None

def classify_content(url: str, title: str, text: str) -> Tuple[str, str, int]:
    """
    Returns (section, page_type, priority) based on multi-signal weighted scoring across URL, Title, and Content.
    """
    url_low = (url or "").lower()
    title_low = (title or "").lower()
    text_low = (text[:4000] or "").lower()

    best_section = "General"
    best_page_type = "GENERAL"
    best_priority = 40
    highest_score = 0

    for section_name, keywords, priority, page_type in SECTION_RULES:
        if not keywords:
            continue
        score = 0
        for kw in keywords:
            kw_low = kw.lower()
            if kw_low in url_low:
                score += 15  # High signal if present in URL path or hostname
            if kw_low in title_low:
                score += 10  # High signal in title/headings
            if kw_low in text_low:
                score += 2   # Context signal in body text

        if score > highest_score:
            highest_score = score
            best_section = section_name
            best_page_type = page_type
            best_priority = priority

    return best_section, best_page_type, best_priority

def extract_html_page(html_content: str or bytes, url: str) -> Dict[str, Any]:
    """Extracts structured text, headings, meta tags, and links from HTML."""
    if isinstance(html_content, bytes):
        soup = BeautifulSoup(html_content, "html.parser", from_encoding="utf-8")
    else:
        soup = BeautifulSoup(html_content, "html.parser")

    # Remove non-content elements
    for tag in soup(["script", "style", "noscript", "svg", "canvas", "template"]):
        tag.decompose()

    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    
    meta_desc = ""
    for meta in soup.find_all("meta"):
        name = (meta.get("name") or meta.get("property") or "").lower()
        if name in ("description", "og:description", "twitter:description"):
            meta_desc = meta.get("content", "").strip()
            break

    # Extract headings
    headings = []
    for h in soup.find_all(["h1", "h2", "h3", "h4"]):
        h_text = h.get_text(" ", strip=True)
        if h_text and len(h_text) > 2:
            headings.append(f"{h.name.upper()}: {h_text}")

    # Extract all discovered hyperlinks
    discovered_links = []
    for a in soup.find_all("a", href=True):
        href = a["href"].strip()
        link_text = clean_extracted_text(a.get_text(" ", strip=True))
        if href and not href.startswith("javascript:") and not href.startswith("mailto:") and not href.startswith("tel:"):
            norm_link = normalize_url(href, base_url=url)
            if norm_link:
                discovered_links.append({"url": norm_link, "text": link_text})

    # Main content container
    main_el = soup.find("main") or soup.find("article") or soup.find("div", class_=re.compile(r'content|body|post|notice', re.I)) or soup.body or soup
    body_text = clean_extracted_text(main_el.get_text("\n", strip=True))

    lang = detect_language(title + " " + body_text)
    pub_date = extract_publication_date(title + "\n" + body_text[:4000])

    section, page_type, priority = classify_content(url, title, body_text)

    return {
        "title": title or "National University Academic Portal",
        "description": meta_desc,
        "clean_text": body_text,
        "headings": headings,
        "links": discovered_links,
        "language": lang,
        "published_date": pub_date,
        "section": section,
        "page_type": page_type,
        "priority": priority,
        "content_hash": compute_sha256(body_text)
    }

def extract_document_file(file_bytes: bytes, url: str, content_type: Optional[str] = None) -> Dict[str, Any]:
    """Extracts text from PDF, DOCX, XLSX, CSV, TXT."""
    norm_url = url.lower()
    content_type = (content_type or "").lower()
    
    extracted_text = ""
    page_count = 1
    doc_type = "UNKNOWN"

    if ".pdf" in norm_url or "application/pdf" in content_type:
        doc_type = "PDF"
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            page_count = len(reader.pages)
            parts = []
            for idx, page in enumerate(reader.pages):
                try:
                    txt = page.extract_text() or ""
                    if txt.strip():
                        parts.append(f"[PDF Page {idx + 1}]\n{txt.strip()}")
                except Exception as pe:
                    logger.warning(f"Error reading page {idx + 1} of {url}: {pe}")
            extracted_text = "\n\n".join(parts)
        except Exception as e:
            logger.warning(f"PDF extraction error on {url}: {e}")
            extracted_text = f"[PDF Notice: {Path(urlparse(url).path).name}]"

    elif ".docx" in norm_url or "wordprocessingml" in content_type:
        doc_type = "DOCX"
        if DocxDocument:
            try:
                doc = DocxDocument(io.BytesIO(file_bytes))
                parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                for tbl in doc.tables:
                    for r in tbl.rows:
                        parts.append(" | ".join(c.text.strip() for c in r.cells))
                extracted_text = "\n".join(parts)
            except Exception as e:
                logger.warning(f"DOCX extraction error on {url}: {e}")

    elif ".xlsx" in norm_url or "spreadsheetml" in content_type:
        doc_type = "XLSX"
        if load_workbook:
            try:
                wb = load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
                parts = []
                for s in wb.worksheets:
                    parts.append(f"[Sheet: {s.title}]")
                    for row in s.iter_rows(values_only=True):
                        vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
                        if vals:
                            parts.append(" | ".join(vals))
                extracted_text = "\n".join(parts)
            except Exception as e:
                logger.warning(f"XLSX extraction error on {url}: {e}")

    elif ".csv" in norm_url or "text/csv" in content_type:
        doc_type = "CSV"
        try:
            decoded = file_bytes.decode("utf-8-sig", errors="replace")
            extracted_text = "\n".join(" | ".join(row) for row in csv.reader(io.StringIO(decoded)))
        except Exception as e:
            logger.warning(f"CSV extraction error on {url}: {e}")

    elif ".txt" in norm_url or content_type.startswith("text/plain"):
        doc_type = "TXT"
        extracted_text = file_bytes.decode("utf-8-sig", errors="replace")

    clean_text = clean_extracted_text(extracted_text)
    file_name = Path(urlparse(url).path).name or "document.pdf"
    pub_date = extract_publication_date(clean_text[:3000])

    section, page_type, priority = classify_content(url, file_name, clean_text)

    return {
        "file_name": file_name,
        "document_type": doc_type,
        "file_size": len(file_bytes),
        "page_count": page_count,
        "extracted_text": clean_text,
        "content_hash": compute_sha256(clean_text or file_bytes),
        "published_date": pub_date,
        "section": section if section != "General" else "Documents",
        "page_type": page_type if page_type != "GENERAL" else "DOCUMENT",
        "priority": priority
    }
