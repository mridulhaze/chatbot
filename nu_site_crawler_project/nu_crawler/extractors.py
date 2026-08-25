import csv
import io
from pathlib import Path
from bs4 import BeautifulSoup
from pypdf import PdfReader
from docx import Document
from openpyxl import load_workbook
from .utils import clean_text, language_of, first_date

def extract_html(html, url):
    if isinstance(html, bytes):
        soup = BeautifulSoup(html, "html.parser", from_encoding="utf-8")
    else:
        soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "noscript", "svg", "canvas", "template"]):
        tag.decompose()

    title = soup.title.get_text(" ", strip=True) if soup.title else ""
    meta = {}
    for m in soup.find_all("meta"):
        key = m.get("name") or m.get("property")
        value = m.get("content")
        if key and value:
            meta[key.lower()] = value.strip()

    headings = []
    for h in soup.find_all(["h1", "h2", "h3", "h4"]):
        t = clean_text(h.get_text(" ", strip=True))
        if t:
            headings.append({"level": h.name, "text": t})

    links = []
    for a in soup.find_all("a", href=True):
        links.append({
            "href": a.get("href", "").strip(),
            "text": clean_text(a.get_text(" ", strip=True))
        })

    main = soup.find("main") or soup.find("article") or soup.find("body") or soup
    text = clean_text(main.get_text("\n", strip=True))
    return {
        "title": title,
        "meta": meta,
        "headings": headings,
        "links": links,
        "text": text,
        "language": language_of(text),
        "published_date": first_date(title + "\n" + text[:5000]),
    }

def extract_pdf(data):
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for i, page in enumerate(reader.pages):
        try:
            txt = page.extract_text() or ""
        except Exception as e:
            txt = f"[PDF extraction error page {i+1}: {e}]"
        if txt.strip():
            parts.append(f"[PDF PAGE {i+1}]\n{clean_text(txt)}")
    text = "\n\n".join(parts)
    return {"page_count": len(reader.pages), "text": text, "language": language_of(text)}

def extract_docx(data):
    doc = Document(io.BytesIO(data))
    parts = [clean_text(p.text) for p in doc.paragraphs if clean_text(p.text)]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(clean_text(c.text) for c in row.cells))
    text = "\n".join(parts)
    return {"text": text, "language": language_of(text)}

def extract_xlsx(data):
    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[SHEET: {ws.title}]")
        for row in ws.iter_rows(values_only=True):
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                parts.append(" | ".join(values))
    text = "\n".join(parts)
    return {"text": clean_text(text), "language": language_of(text), "sheets": wb.sheetnames}

def extract_csv(data):
    raw = data.decode("utf-8-sig", errors="replace")
    text = clean_text("\n".join(" | ".join(row) for row in csv.reader(io.StringIO(raw))))
    return {"text": text, "language": language_of(text)}

def extract_txt(data):
    text = clean_text(data.decode("utf-8-sig", errors="replace"))
    return {"text": text, "language": language_of(text)}

def extract_document(data, content_type, suffix):
    ct = (content_type or "").lower()
    if suffix == ".pdf" or "application/pdf" in ct:
        return extract_pdf(data)
    if suffix == ".docx" or "wordprocessingml.document" in ct:
        return extract_docx(data)
    if suffix == ".xlsx" or "spreadsheetml" in ct:
        return extract_xlsx(data)
    if suffix == ".csv" or "text/csv" in ct:
        return extract_csv(data)
    if suffix == ".txt" or ct.startswith("text/plain"):
        return extract_txt(data)
    return {"text": "", "language": "unknown", "unsupported": True}
